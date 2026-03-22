"""
수집된 판례/조정례를 Word(.docx) 파일과 HTML 링크 파일로 출력
기준 형식: 250817_2Q판례정리.txt
"""

import re
from datetime import datetime
from pathlib import Path
from urllib.parse import urlparse, parse_qs

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from summarizer import parse_summary_sections


SOURCE_ORDER = ["대법원", "금감원", "소비자원"]
SOURCE_FULL_NAMES = {
    "대법원": "대법원 판례",
    "금감원": "금융감독원 분쟁조정사례",
    "소비자원": "한국소비자원 분쟁조정결정례",
}


def generate_filename_prefix(year: int, quarter: int) -> str:
    """파일명 접두어 생성: YYMMDD_NQ"""
    today = datetime.now().strftime("%y%m%d")
    return f"{today}_{quarter}Q"


def write_docx(cases: list[dict], year: int, quarter: int, output_dir: str = ".") -> str:
    """
    Word 파일 생성
    반환: 생성된 파일 경로
    """
    prefix = generate_filename_prefix(year, quarter)
    filename = f"{prefix}판례정리.docx"
    filepath = Path(output_dir) / filename

    doc = Document()
    _setup_styles(doc)

    # 실패 사이트 목록 수집
    failed_sources = _get_failed_sources(cases)

    # 커버 헤더
    _add_document_header(doc, year, quarter, failed_sources)

    # 소스별로 그룹화하여 출력
    cases_by_source = _group_by_source(cases)

    for source in SOURCE_ORDER:
        source_cases = cases_by_source.get(source, [])
        if not source_cases:
            continue

        # 소스 섹션 헤더
        _add_source_header(doc, source)

        for case in source_cases:
            _add_case_entry(doc, case)

    doc.save(str(filepath))
    print(f"[출력] Word 파일 저장: {filepath}")
    return str(filepath)


def write_html(cases: list[dict], year: int, quarter: int, output_dir: str = ".") -> str:
    """
    원문 링크 HTML 파일 생성
    반환: 생성된 파일 경로
    """
    prefix = generate_filename_prefix(year, quarter)
    filename = f"{prefix}_원문링크.html"
    filepath = Path(output_dir) / filename

    cases_by_source = _group_by_source(cases)
    failed_sources = _get_failed_sources(cases)

    html_content = _build_html(year, quarter, cases_by_source, failed_sources, SOURCE_ORDER)

    filepath.write_text(html_content, encoding="utf-8")
    print(f"[출력] HTML 파일 저장: {filepath}")
    return str(filepath)


# ── 내부 함수 ──────────────────────────────────────────────────────────────────


def _get_failed_sources(cases: list[dict]) -> list[str]:
    """수집에 실패한 소스 목록 반환"""
    collected = {c.get("source", "") for c in cases}
    return [s for s in SOURCE_ORDER if s not in collected]


def _group_by_source(cases: list[dict]) -> dict[str, list]:
    groups: dict[str, list] = {}
    for case in cases:
        source = case.get("source", "기타")
        groups.setdefault(source, []).append(case)
    return groups


def _setup_styles(doc: Document):
    """기본 스타일 설정"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(10)


def _add_document_header(doc: Document, year: int, quarter: int, failed_sources: list[str]):
    """문서 제목 및 메타 정보"""
    title = doc.add_heading(f"{year}년 {quarter}분기 판례/분쟁조정사례 정리", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 생성일
    info = doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if failed_sources:
        failed_para = doc.add_paragraph()
        run = failed_para.add_run(f"⚠️ 수집 실패 사이트: {', '.join(failed_sources)}")
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        failed_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # 빈 줄


def _add_source_header(doc: Document, source: str):
    """소스 구분 헤더"""
    full_name = SOURCE_FULL_NAMES.get(source, source)
    heading = doc.add_heading(full_name, level=2)
    # 구분선
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4472C4")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_case_entry(doc: Document, case: dict):
    """케이스 항목 추가 (4섹션 구조)"""
    title = case.get("title", "제목 없음")
    date = case.get("date", "")
    url = case.get("url", "")
    summary_text = case.get("summary", "")

    # 케이스 제목
    case_heading = doc.add_heading(title, level=3)

    # 날짜 + 원문 링크 (소제목처럼)
    meta = doc.add_paragraph()
    meta.add_run(f"날짜: {date}").font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    if url:
        meta.add_run("  |  ")
        run = meta.add_run("원문 링크")
        run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
        run.font.underline = True
        # 하이퍼링크 추가
        _add_hyperlink(meta, url, "원문 링크")

    # 4섹션 내용
    if summary_text:
        sections = parse_summary_sections(summary_text)
        section_labels = [
            ("사건_요약", "사건 요약"),
            ("사실_관계", "사실 관계"),
            ("판시_사항", "판시 사항"),
            ("활용_방안", "활용 방안"),
        ]
        for i, (key, label) in enumerate(section_labels, 1):
            content = sections.get(key, "")
            # 섹션 제목
            sec_para = doc.add_paragraph()
            run = sec_para.add_run(f"{i}. {label}")
            run.bold = True
            run.font.size = Pt(10)
            # 섹션 내용
            if content:
                doc.add_paragraph(content)
            else:
                doc.add_paragraph("[내용 없음]")
    else:
        doc.add_paragraph("[요약 생성 실패]")

    # 케이스 구분 여백
    doc.add_paragraph()


def _add_hyperlink(paragraph, url: str, text: str):
    """Word 문서에 하이퍼링크 추가"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


def _build_html(
    year: int, quarter: int,
    cases_by_source: dict,
    failed_sources: list[str],
    source_order: list[str]
) -> str:
    """HTML 파일 내용 생성"""
    today = datetime.now().strftime("%Y년 %m월 %d일")

    # 탭 버튼 HTML
    tab_buttons = ""
    tab_contents = ""

    for i, source in enumerate(source_order):
        source_cases = cases_by_source.get(source, [])
        full_name = SOURCE_FULL_NAMES.get(source, source)
        active = "active" if i == 0 else ""

        tab_buttons += f'<button class="tab-btn {active}" onclick="showTab(\'{source}\')" id="btn-{source}">{full_name} ({len(source_cases)}건)</button>\n'

        items_html = ""
        if not source_cases:
            items_html = '<div class="no-data">수집된 데이터가 없습니다.</div>'
        else:
            for case in source_cases:
                url = case.get("url", "")
                title = case.get("title", "제목 없음")
                date = case.get("date", "")
                if source == "소비자원" and url:
                    qs = parse_qs(urlparse(url).query)
                    seq = qs.get("seq", [""])[0]
                    brd_id = qs.get("brdId", ["00000007"])[0]
                    link_html = (
                        f'<form method="POST" action="https://www.kca.go.kr/odr/cm/cm/boardsDtl.do" '
                        f'target="_blank" style="display:inline;margin:0;">'
                        f'<input type="hidden" name="brdId" value="{brd_id}">'
                        f'<input type="hidden" name="seq" value="{seq}">'
                        f'<input type="hidden" name="dataStts" value="Y">'
                        f'<button type="submit" class="case-link post-btn">{title}</button>'
                        f'</form>'
                    )
                elif url:
                    link_html = f'<a href="{url}" target="_blank" class="case-link">{title}</a>'
                else:
                    link_html = f'<span>{title}</span>'
                items_html += f"""
                <div class="case-item">
                    <div class="case-title">{link_html}</div>
                    <div class="case-date">📅 {date}</div>
                </div>"""

        display = "block" if i == 0 else "none"
        tab_contents += f"""
        <div id="tab-{source}" class="tab-content" style="display:{display};">
            {items_html}
        </div>"""

    failed_warning = ""
    if failed_sources:
        failed_warning = f"""
        <div class="warning">
            ⚠️ 다음 사이트에서 데이터 수집에 실패했습니다: <strong>{", ".join(failed_sources)}</strong>
        </div>"""

    return f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{year}년 {quarter}분기 판례 원문 링크</title>
    <style>
        body {{
            font-family: 'Malgun Gothic', '맑은 고딕', sans-serif;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background: #f5f5f5;
            color: #333;
        }}
        h1 {{
            color: #1a3a6b;
            border-bottom: 3px solid #4472C4;
            padding-bottom: 10px;
        }}
        .meta {{
            color: #666;
            margin-bottom: 20px;
        }}
        .warning {{
            background: #fff3cd;
            border: 1px solid #ffc107;
            border-radius: 4px;
            padding: 12px;
            margin-bottom: 20px;
        }}
        .tabs {{
            display: flex;
            gap: 4px;
            margin-bottom: 20px;
            flex-wrap: wrap;
        }}
        .tab-btn {{
            padding: 10px 18px;
            border: 1px solid #4472C4;
            background: white;
            color: #4472C4;
            cursor: pointer;
            border-radius: 4px 4px 0 0;
            font-size: 14px;
            font-family: inherit;
        }}
        .tab-btn.active {{
            background: #4472C4;
            color: white;
        }}
        .tab-content {{
            background: white;
            border: 1px solid #ddd;
            border-radius: 0 4px 4px 4px;
            padding: 20px;
        }}
        .case-item {{
            border-bottom: 1px solid #eee;
            padding: 14px 0;
        }}
        .case-item:last-child {{ border-bottom: none; }}
        .case-title {{
            font-size: 15px;
            margin-bottom: 4px;
        }}
        .case-link {{
            color: #0056b3;
            text-decoration: none;
        }}
        .case-link:hover {{ text-decoration: underline; }}
        .post-btn {{
            background: none;
            border: none;
            padding: 0;
            cursor: pointer;
            font-size: 15px;
            font-family: inherit;
            text-align: left;
        }}
        .case-date {{ color: #888; font-size: 13px; }}
        .no-data {{ color: #888; padding: 20px; text-align: center; }}
        .total {{
            text-align: right;
            color: #666;
            font-size: 13px;
            margin-top: 12px;
        }}
    </style>
</head>
<body>
    <h1>{year}년 {quarter}분기 판례/분쟁조정사례 원문 링크</h1>
    <div class="meta">생성일: {today}</div>
    {failed_warning}
    <div class="tabs">
        {tab_buttons}
    </div>
    {tab_contents}

    <script>
        function showTab(source) {{
            document.querySelectorAll('.tab-content').forEach(el => el.style.display = 'none');
            document.querySelectorAll('.tab-btn').forEach(el => el.classList.remove('active'));
            document.getElementById('tab-' + source).style.display = 'block';
            document.getElementById('btn-' + source).classList.add('active');
        }}
    </script>
</body>
</html>"""
